from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def highlight_paragraphs(file_path, paragraphs_to_highlight):
    # Load the document
    doc = Document(file_path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Check if the paragraph matches the ones to highlight
                    if any(text in paragraph.text for text in paragraphs_to_highlight):
                        # Highlight the paragraph
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


    # Save the modified document
    doc.save('highlighted_doc.docx')

# Example usage
file_path = './1.docx'
paragraphs_to_highlight = ['This statement']

highlight_paragraphs(file_path, paragraphs_to_highlight)