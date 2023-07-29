from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm
import docx
from docx.text.paragraph import Paragraph
import os
import docx2txt
from win32com import client as wc

class DocumentReader:
    """
    Return a |Document| object loaded from *docx*.
    """
    def __init__(self, file_path):
        self.file_path = file_path

    def read_runs_from_paragraph(self, paragraph):
        runs = paragraph.runs
        return [run.text for run in runs]

    def iter_paragraphs(self, parent):
        """
        Generate a reference to each paragraph within *parent*, in document order.
        Each returned value is an instance of Paragraph. *parent* would most
        commonly be a reference to a main Document object, but also works for
        a Cell or Row. Note that each call to iter_paragraphs() returns a new
        generator each time it is called. The same set of Paragraph references
        cannot be iterated over more than once.
        """
        doc = parent if isinstance(parent, docx.Document) else parent._parent
        for child in parent._element.body:
            if isinstance(child, docx.oxml.text.paragraph.CT_P):
                yield Paragraph(child, doc)
            elif isinstance(child, docx.oxml.table.CT_Tbl):
                for row in child.tblPr:
                    for cell in row.tcPr:
                        for grandchild in cell.body:
                            if isinstance(grandchild, docx.oxml.text.paragraph.CT_P):
                                yield Paragraph(grandchild, doc)

    def read_docx(self):
        document = Document(self.file_path)
        all_runs = []

        for paragraph in self.iter_paragraphs(document):
            all_runs.extend(self.read_runs_from_paragraph(paragraph))

        return all_runs

class Reviewer:
    def __init__(self):
        self.en_path = r"C:\Users\sai4w\Downloads\KFS_Paired_document_(1)\KFS Paired document\KFS - CT UK Smaller Companies Fund_E.DOCX"
        self.ch_path = r"C:\Users\sai4w\Downloads\KFS_Paired_document_(1)\KFS Paired document\KFS - CT UK Smaller Companies Fund_C.DOCX"
    def result(self, paragraphs_to_highlight):
        
        # Load the document
        doc = Document(self.en_path)
        print('read doc')
        # Process and modify paragraphs
        for paragraph in doc.paragraphs:
            self.add_paragraph_with_highlight_and_additional_text(paragraph, paragraphs_to_highlight)

        # Process and modify tables
        for table in doc.tables:
            self.process_table(table, paragraphs_to_highlight)

        # Save the modified document
        print("hereee")
        doc.save('highlighted_doc.docx')
    def add_error_message(self, run,message):
        run.add_break()
        run.add_text(message)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW  
        run.add_break()
    def add_chinese_paragraph(self, run, chinese_text):
        run.add_break()
        run.add_text(chinese_text)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW  
        run.add_break()
    def add_paragraph_with_highlight_and_additional_text(self, paragraph, paragraphs_to_highlight):
        for target_para_info in paragraphs_to_highlight:
            target_paragraph = target_para_info["english"]
            chinese_paragraph = target_para_info["chinese"]
            if target_paragraph != '' and target_paragraph in paragraph.text:
                # Split the paragraph text into parts around the target paragraph
                self.highlight_words_in_paragraph(paragraph, target_paragraph)
                additional_text = '\n Errror in text' +self.build_additional_message(target_para_info)+'\n'
                paragraph.paragraph_format.keep_together = True
                run =paragraph.add_run()
                run.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                self.add_error_message(run, additional_text)
                run =paragraph.add_run()
                run.font.size = Pt(12)
                run.font.name = 'Microsoft YaHei'
                self.add_chinese_paragraph(run, chinese_paragraph)
    def process_table(self, table, paragraphs_to_highlight):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.add_paragraph_with_highlight_and_additional_text(paragraph, paragraphs_to_highlight)
    def create_highlighted_run(self, run, texts, target_paragraph):

        for text in texts:
            if text == target_paragraph:
                run.add_text(text)
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW            
    def build_additional_message(self, paragraph_object):
        message = ""
        if 'numbers' in paragraph_object and paragraph_object["numbers"] == 0:
            message += error_messages["numbers_error"]
        if  "para_score" in paragraph_object and paragraph_object["para_score"] ==2 :
            message += "\n" + error_messages["could_be_inaccurate_translation_2"]
        if "para_score" in paragraph_object and paragraph_object["para_score"]==3:
            message += "\n" + error_messages["inaccurate_translation_or_missing_3"]
        return message
    def highlight_words_in_paragraph(self, paragraph, word_to_highlight, highlight_color=WD_COLOR_INDEX.YELLOW):
        text = paragraph.text
        start_index = text.find(word_to_highlight)
        
        # Only proceed if the word is found in the paragraph
        if start_index != -1:
            end_index = start_index + len(word_to_highlight)
            
            # Clear any existing highlighting in the paragraph
            for run in paragraph.runs:
                run.font.highlight_color = None
                
            # Apply highlighting to the specific word
            for run in paragraph.runs:
                if start_index >= len(run.text):
                    start_index -= len(run.text)
                    end_index -= len(run.text)
                else:
                    if end_index <= len(run.text):
                        run.start = start_index
                        run.end = end_index
                        run.font.highlight_color = highlight_color
                        break
                    else:
                        run.start = start_index
                        run.end = len(run.text)
                        run.font.highlight_color = highlight_color
                        start_index = 0
                        end_index -= len(run.text)

error_messages = {
    "numbers_error":"Numbers are not matching",
    "could_be_inaccurate_translation_2":"Could be inaccurate translation",
    "inaccurate_translation_or_missing_3":"Could be inaccurate translation",
}

x = Reviewer()
paragraphs = [{'para_score': 2, 'english': 'Threadneedle Portfolio Services Hong Kong Limited', 'chinese': '產品資料概要'}, {'para_score': 2, 'english': 'PRODUCT KEY FACTS', 'chinese': '天利（英國）投資基金'}, {'para_score': 2, 'english': 'Columbia Threadneedle Investment Funds (UK) ICVC', 'chinese': '天利英國小型公司基金'}, {'para_score': 2, 'english': 'CT UK Smaller Companies Fund', 'chinese': '2022年7月[*]日'}, {'para_score': 3, 'english': '22 September 2022', 'chinese': '發行人', 'numbers': 0}, {'para_score': 2, 'english': 'Columbia Threadneedle Investment Funds (UK) ICVC – CT UK Smaller Companies Fund', 'chinese': '天利投資管理香港有限公司'}, {'para_score': 2, 'english': 'This statement is a part of the offering document.', 'chinese': '本概要是銷售文件的一部分。'}, {'para_score': 3, 'english': 'Quick facts', 'chinese': '資料便覽'}, {'para_score': 2, 'english': 'Threadneedle Investment Services Limited', 'chinese': 'Threadneedle Investment Services Limited'}, {'para_score': 2, 'english': 'Citibank UK Limited', 'chinese': 'Citibank UK Limited'}, {'para_score': 2, 'english': 'Citibank N.A.', 'chinese': 'Citibank N.A.'}, {'para_score': 2, 'english': 'Class 1', 'chinese': '第1 類'}, {'para_score': 2, 'english': 'Daily', 'chinese': '每日', 'numbers': 0}, {'para_score': 2, 'english': 'Income Shares', 'chinese': '收入股份'}, {'para_score': 2, 'english': 'Class 1', 'chinese': '第1 類'}, {'para_score': 2, 'english': 'GBP 2,000 initial, GBP 1,000 additional', 'chinese': '首次投資額為2,000 英鎊，其後認購額為1,000 英鎊'}, {'para_score': 2, 'english': 'What is this product?', 'chinese': '本基金是甚麼產品？'}, {'para_score': 2, 'english': 'Investment', 'chinese': '投資風險'}, {'para_score': 3, 'english': 'Volatility', 'chinese': '波動性風險'}, {'para_score': 2, 'english': 'Smaller Companies', 'chinese': '小型公司風險'}, {'para_score': 2, 'english': 'Geographical Concentration', 'chinese': '地域集中性風險'}, {'para_score': 2, 'english': 'Currency', 'chinese': '貨幣風險'}, {'para_score': 2, 'english': 'Investor Currency', 'chinese': '投資者貨幣風險'}, {'para_score': 2, 'english': 'Hedge / Basis', 'chinese': '對沖╱基準風險'}, {'para_score': 2, 'english': 'Is there any guarantee?', 'chinese': '本基金是否提供任何保證？'}, {'para_score': 2, 'english': 'This Fund does not have any guarantees. You may not get back the full amount of money you invest.', 'chinese': '本基金並不提供任何保證。閣下未必能全數收回投資本金。'}, {'para_score': 2, 'english': 'Charges which may be payable by you', 'chinese': '閣下或須繳付的收費'}, {'para_score': 2, 'english': 'You may have to pay the following fees when dealing in the shares of the Fund.', 'chinese': '閣下就本基金股份交易須繳付以下費用。'}, {'para_score': 2, 'english': 'What you pay', 'chinese': '您須繳付'}, {'para_score': 2, 'english': 'Subscription Fee', 'chinese': '認購費'}, {'para_score': 2, 'english': '(Initial Charge)', 'chinese': '（認購費用）'}, {'para_score': 2, 'english': 'Class 1', 'chinese': '第1 類'}, {'para_score': 2, 'english': 'Up to the prevailing initial charge for the class of shares acquired', 'chinese': '最高為所認購股份所屬類別的首次認購費金額'}, {'para_score': 2, 'english': '(Redemption Charge)', 'chinese': '（贖回費用）'}, {'para_score': 2, 'english': 'Annual rate', 'chinese': '每年收費率'}, {'para_score': 2, 'english': 'Class 1', 'chinese': '第1 類'}, {'para_score': 2, 'english': 'Custodian fee', 'chinese': '保管費'}, {'para_score': 2, 'english': 'Performance fee', 'chinese': '表現費'}, {'para_score': 2, 'english': '(Ongoing registration and general expenses payable to the ACD)', 'chinese': '（持續向受權公司董事支付的登記費用及一般費用）'}, {'para_score': 2, 'english': 'Class 1', 'chinese': '第1 類'}, {'para_score': 2, 'english': 'You may have to pay other fees when dealing in the shares of the Fund.', 'chinese': '本基金股份交易或須繳付其他費用。'}, {'para_score': 2, 'english': 'Additional Information', 'chinese': '其他資料'}, {'para_score': 2, 'english': 'Important', 'chinese': '重要提示', 'numbers': 0}, {'para_score': 2, 'english': 'main\\vwwong\\10089492_1.doc', 'chinese': 'CT-00042201/v1.0'}, {'para_score': 2, 'english': 'Error! Unknown document property name.CT-00020306/v1.01', 'chinese': '2'}, {'para_score': 2, 'english': '', 'chinese': '3', 'numbers': 0}, {'para_score': 3, 'english': '', 'chinese': 'MAIN\\HOAMY\\36635956_1.docx', 'numbers': 0}, {'para_score': 3, 'english': '', 'chinese': 'CT-00042201/v1.0', 'numbers': 0}, {'para_score': 3, 'english': '', 'chinese': '1', 'numbers': 0}, {'para_score': 3, 'english': '', 'chinese': 'MAIN\\HOAMY\\36635956_1.docx', 'numbers': 0}]
x.result(paragraphs)





def extract_text_from_docx(path):
    temp = docx2txt.process(path)
    text = [line.replace('\t', ' ') for line in temp.split('\n') if line]
    final_text = ' '.join(text)
    return final_text

def extract_text_from_doc(doc_path):
    w = wc.Dispatch('Word.Application')
    doc = w.Documents.Open(doc_path)
    doc.SaveAs(save_file_name, 16)
    doc.Close()
    w.Quit()
    joinedPath = os.path.join(save_file_name)
    text = extract_text_from_docx(joinedPath)
    return text

def extract_text(file_path, extension):
    text = ''
    if extension in ['.docx', '.DOCX'] :
        text = extract_text_from_docx(file_path)
    elif extension == ['.doc', '.DOC'] :
        text = extract_text_from_doc(file_path)
    return text

if __name__ == "__main__":
    print("Running")
    file_path = r"C:\Users\sai4w\Downloads\KFS_Paired_document_(1)\KFS Paired document\KFS - CT UK Smaller Companies Fund_E.DOCX"
    save_file_name = "resultat_docx.docx"
    extension = os.path.splitext(file_path)[1]
    final_text = extract_text(file_path, extension)
    print(final_text)
