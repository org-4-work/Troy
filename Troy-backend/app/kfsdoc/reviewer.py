import os
import uuid
import json
from flask import jsonify
import traceback
import pandas as pd
from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from app.kfsdoc.kfsdoc import KFSDocReviewer


MISSING_TRANSLATION_ACCURACY = 1.1
UNACURATE_TRANSLATION_ACCURACY = 1.4

class Reviewer:
    def __init__(self, en_path, ch_path):
        self.en_path = en_path
        self.ch_path = ch_path

    def start_review(self):
        """
        Start the review process by initializing KFSDocReviewer and processing the results.
        """
        
        print("Starting review")
        kfc_document = KFSDocReviewer(self.en_path, self.ch_path)
        print("Initializing KFSDocReviewer")
        kfc_document.initialize()
        print("Processing results")
        global english_sentences_array
        global ai_translated_array
        global chinese_array
        global translation_accuracy_array
        global number_matching_scores_array


        print("Number : 1")
        english_sentences_array, ai_translated_array, chinese_array, translation_accuracy_array, number_matching_scores_array = kfc_document.treat()
        print("Number : 2")
        # number_matching_scores_array = [
        #     eval(item) for item in number_matching_scores_array]
        print("Number : 3")
        english_sentences_array, ai_translated_array, chinese_array, translation_accuracy_array, number_matching_scores_array = remove_wrong_lines(
        english_sentences_array, ai_translated_array, chinese_array, translation_accuracy_array, number_matching_scores_array)
        print('generating')
        highlited_file = generate_result(english_sentences_array, self.en_path)
        print('successfully')

        return jsonify({"docx_file": highlited_file})    

def remove_wrong_lines(english_sentences_array, ai_translated_array, chinese_array, translation_accuracy_array, number_matching_scores_array):
    list_of_wrong_lines = [
        '.docx', "Unknown document property name.", "Error!"]
    indexes_to_keep = [index for index, item in enumerate(english_sentences_array) if (
        isinstance(item, str) and (item != '' or (not item.isdigit()) or not (item in list_of_wrong_lines)))]

    english_sentences_array = [english_sentences_array[index]
                               for index in indexes_to_keep]
    ai_translated_array = [ai_translated_array[index]
                           for index in indexes_to_keep]
    chinese_array = [chinese_array[index] for index in indexes_to_keep]
    translation_accuracy_array = [
        translation_accuracy_array[index] for index in indexes_to_keep]
    number_matching_scores_array = [
        number_matching_scores_array[index] for index in indexes_to_keep]

    return english_sentences_array, ai_translated_array, chinese_array, translation_accuracy_array, number_matching_scores_array




def highlight_words_in_paragraph(paragraph, word_to_highlight, highlight_color=WD_COLOR_INDEX.YELLOW):
    text = paragraph.text
    start_index = text.find(word_to_highlight)

    # Only proceed if the word is found in the paragraph
    if start_index == -1:
        return
    end_index = start_index + len(word_to_highlight)

    # Clear any existing highlighting in the paragraph

    # Apply highlighting to the specific word
    for run in paragraph.runs:
        if start_index >= len(run.text):
            start_index -= len(run.text)
            end_index -= len(run.text)
        else:
            if end_index <= len(run.text):
                run.start = start_index
                run.end = end_index
                run.font.highlight_color = highlight_color
                break
            else:
                run.start = start_index
                run.end = len(run.text)
                run.font.highlight_color = highlight_color
                start_index = 0
                end_index -= len(run.text)


def highlight_number_in_run(run, number_to_highlight, highlight_color=WD_COLOR_INDEX.RED):
    text = run.text
    start_index = text.find(number_to_highlight)

    # Only proceed if the word is found in the paragraph
    if start_index == -1:
        return
    end_index = start_index + len(number_to_highlight)
    if start_index >= len(run.text):
        start_index -= len(run.text)
        end_index -= len(run.text)
    else:
        if end_index <= len(run.text):
            run.start = start_index
            run.end = end_index
            run.font.highlight_color = highlight_color
        else:
            run.start = start_index
            run.end = len(run.text)
            run.font.highlight_color = highlight_color
            start_index = 0
            end_index -= len(run.text)


def process_table(table, paragraphs_to_highlight):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                add_paragraph_with_highlight_and_additional_text(
                    paragraph, paragraphs_to_highlight)
            for nested_table in cell.tables:
                process_table(nested_table, paragraphs_to_highlight)


def process_headers_footers(section, paragraphs_to_highlight):
    for paragraph in section.header.paragraphs:
        add_paragraph_with_highlight_and_additional_text(
            paragraph, paragraphs_to_highlight)
    for table in section.header.tables:
        process_table(table, paragraphs_to_highlight)

    for paragraph in section.footer.paragraphs:
        add_paragraph_with_highlight_and_additional_text(
            paragraph, paragraphs_to_highlight)
    for table in section.footer.tables:
        process_table(table, paragraphs_to_highlight)


def generate_result(paragraphs_to_highlight, path):
    # Load the document
    doc = Document(path)
    print('read doc')
    unique_filename = str(f"output/{uuid.uuid4()} _ highlighted.docx")

    # Process and modify paragraphs in the main body
    for paragraph in doc.paragraphs:
        add_paragraph_with_highlight_and_additional_text(
            paragraph, paragraphs_to_highlight)

    # Process and modify tables in the main body
    for table in doc.tables:
        process_table(table, paragraphs_to_highlight)

    # Process and modify headers and footers for each section in the document
    for section in doc.sections:
        process_headers_footers(section, paragraphs_to_highlight)
    doc.save(path)
    return path



def write_additional_text(run, text):
    run.add_break()
    run.add_text(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.add_break()
    return run


def add_paragraph_with_highlight_and_additional_text(paragraph, paragraphs_to_highlight):
    for index, target_paragraph in enumerate(paragraphs_to_highlight):
        # target_paragraph = target_para_info["english"]
        # chinese_paragraph = target_para_info["chinese"]
        target_paragraph = str(target_paragraph)
        if target_paragraph != '' and target_paragraph in paragraph.text:
            # Split the paragraph text into parts around the target paragraph
            if (translation_accuracy_array[index] < MISSING_TRANSLATION_ACCURACY):
                highlight_words_in_paragraph(
                    paragraph, target_paragraph, WD_COLOR_INDEX.RED)
            elif (translation_accuracy_array[index] < UNACURATE_TRANSLATION_ACCURACY):
                highlight_words_in_paragraph(
                    paragraph, target_paragraph, WD_COLOR_INDEX.YELLOW)
                run = paragraph.add_run()
                run.font.size = Pt(12)
                run.font.name = 'Microsoft YaHei'
                run.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = write_additional_text(run, chinese_array[index])
                for nb in number_matching_scores_array[index]['missing_numbers']:
                    highlight_words_in_paragraph(
                        paragraph, str(nb), WD_COLOR_INDEX.BLUE)

                for nb in number_matching_scores_array[index]['incorrect_numbers']:
                    highlight_number_in_run(run, str(nb), WD_COLOR_INDEX.BLUE)
            
